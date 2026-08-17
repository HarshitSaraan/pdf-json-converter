import re
import os
import json
import pypdf
import docx

def extract_raw_text(file_path: str) -> str:
    """Extract raw text from PDF or DOCX file."""
    ext = os.path.splitext(file_path)[1].lower()
    full_text = ""

    if ext in ['.docx', '.doc']:
        try:
            doc = docx.Document(file_path)
            lines = []
            for p in doc.paragraphs:
                if p.text.strip():
                    lines.append(p.text.strip())
            for t in doc.tables:
                for row in t.rows:
                    row_txt = ' | '.join([cell.text.strip() for cell in row.cells if cell.text.strip()])
                    if row_txt:
                        lines.append(row_txt)
            full_text = "\n".join(lines)
        except Exception as e:
            print(f"docx extraction failed: {e}")

    if not full_text.strip():
        # Try pypdf (pure Python, lightweight for Vercel serverless)
        try:
            reader = pypdf.PdfReader(file_path)
            pages_text = [p.extract_text() or "" for p in reader.pages]
            full_text = "\n".join(pages_text)
        except Exception as e:
            print(f"pypdf extraction failed: {e}")

    if not full_text.strip():
        try:
            import fitz
            doc = fitz.open(file_path)
            pages_text = [page.get_text("text") or "" for page in doc]
            full_text = "\n".join(pages_text)
        except Exception as e:
            print(f"fitz extraction failed: {e}")

    if not full_text.strip():
        try:
            import pdfplumber
            with pdfplumber.open(file_path) as pdf:
                pages_text = [p.extract_text() or "" for p in pdf.pages]
                full_text = "\n".join(pages_text)
        except Exception as e:
            print(f"pdfplumber extraction failed: {e}")

    return full_text

def format_math_latex(text: str) -> str:
    r"""
    Ensure math expressions and text are properly formatted according to strict rules:
    1. Inline-safe LaTeX only ($...$), convert \[ ... \] or $$ ... $$ to $...$.
    2. Don't write Step 1, Step 2... instead just line by line.
    3. Don't use bold letters (remove ** and __).
    4. Don't use ### or # headings.
    """
    if not text:
        return ""
    
    # Remove markdown headings and '###'
    text = re.sub(r'#+\s*', '', text)

    # Convert display LaTeX \[...\] or $$...$$ to inline $...$
    text = text.replace(r'\[', '$').replace(r'\]', '$')
    text = text.replace('$$', '$')

    # Remove newlines inside $...$ to ensure inline-safety (same line)
    def inline_latex(match):
        latex_content = match.group(1).strip()
        latex_content = re.sub(r'\s+', ' ', latex_content)
        return f"${latex_content}$"
    text = re.sub(r'\$([^\$\n]+)\$', inline_latex, text)

    # Remove "Step 1", "Step 2", "Step 1:", "Step 2.", "Step 1 -", etc.
    text = re.sub(r'(?i)\bStep\s*\d+[:\.-]?\s*', '', text)

    # Remove bold syntax (**text** or __text__ or stray ** / __)
    text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
    text = re.sub(r'__(.*?)__', r'\1', text)
    text = text.replace('**', '').replace('__', '')

    # Clean up whitespace
    text = re.sub(r'[ \t]+', ' ', text)
    lines = [line.strip() for line in text.splitlines()]
    text = "\n".join(lines).strip()
    text = re.sub(r'\n{3,}', '\n\n', text)
    return text


OFFICIAL_TAXONOMY = {
    "English": {
        "VA": [
            "RC",
            "Para Completion",
            "Para Jumbles",
            "Sentence Correction",
            "Spellings",
            "Verbal Analogy"
        ],
        "Vocabulary": [
            "Idioms & Phrases",
            "Antonyms",
            "Synonyms",
            "Definition"
        ],
        "Grammar": [
            "Active & Passive Voice",
            "Direct & Indirect Speech",
            "Error",
            "Punctuations",
            "Parts of Speech",
            "Subject–Verb Agreement"
        ]
    },
    "Quants": {
        "Arithmetic": [
            "Averages",
            "Mixtures & Alligation",
            "Percentages",
            "Profit & Loss",
            "Ratio & Proportion",
            "SI/CI",
            "Time & Work",
            "Time–Speed–Distance"
        ],
        "Number System": [
            "Digit properties",
            "Divisibility rules",
            "Factorials",
            "Factorization",
            "Factors/Multiples",
            "HCF/LCM",
            "Integral Solution",
            "Miscellaneous",
            "Remainders",
            "Unit digits"
        ],
        "Algebra": [
            "Binomial Theorem",
            "Matrices & Determinants",
            "Algebraic identities",
            "Functions",
            "Indices & Surds",
            "Inequalities",
            "Linear/Quadratic equations",
            "Maxima & Minima",
            "Modulus",
            "Polynomials",
            "Progressions",
            "Sets"
        ],
        "Geometry & Mensuration": [
            "Area & Perimeter",
            "Circles",
            "Coordinate Geometry",
            "Heights & Distances",
            "Lines & Angles",
            "Polygons",
            "Quadrilaterals",
            "Solids",
            "Triangles",
            "Trigonometry"
        ],
        "Modern Maths": [
            "Binomial Theorem",
            "Logarithm",
            "Matrices & Determinants",
            "P & C",
            "Probability",
            "Set Theory"
        ]
    },
    "LRDI": {
        "Logical Reasoning": [
            "Arrangements",
            "Blood Relations",
            "Clocks & Calendars",
            "Coding-Decoding",
            "Direction Sense",
            "Syllogisms",
            "Series & Analogies",
            "Venn Diagrams"
        ],
        "Data Interpretation": [
            "Bar Charts",
            "Line Graphs",
            "Pie Charts",
            "Tables",
            "Caselets",
            "Data Sufficiency"
        ]
    }
}

ENGLISH_PATTERNS = [
    # Vocabulary
    ("Vocabulary", "Idioms & Phrases", [(r'idiom', 5), (r'phrasal verb', 5), (r'phrase', 3), (r'proverb', 4), (r'went by the book', 5), (r'spill the beans', 5), (r'burn the midnight oil', 5), (r'kick the bucket', 5), (r'apple of discord', 5)]),
    ("Vocabulary", "Synonyms", [(r'synonym', 5), (r'similar in meaning', 5), (r'closest in meaning', 5), (r'same meaning', 4), (r'nearest in meaning', 5)]),
    ("Vocabulary", "Antonyms", [(r'antonym', 5), (r'opposite in meaning', 5), (r'contrary in meaning', 5), (r'furthest in meaning', 5), (r'reverse in meaning', 4)]),
    ("Vocabulary", "Spellings", [(r'correctly spelt', 5), (r'misspelt', 5), (r'spelling', 4), (r'correct spelling', 5), (r'incorrectly spelt', 5)]),
    ("Vocabulary", "Definition", [(r'meaning of', 3), (r'what does .* mean', 3), (r'definition of', 4), (r'word .* means', 3), (r'refers to', 2)]),

    # Grammar
    ("Grammar", "Active & Passive Voice", [(r'active voice', 5), (r'passive voice', 5), (r'change the voice', 5), (r'convert into passive', 5), (r'convert into active', 5)]),
    ("Grammar", "Direct & Indirect Speech", [(r'direct speech', 5), (r'indirect speech', 5), (r'reported speech', 5), (r'convert into indirect', 5), (r'said that', 3)]),
    ("Grammar", "Error", [(r'spot the error', 5), (r'error spotting', 5), (r'find the error', 5), (r'grammatically incorrect', 5), (r'error in part', 5), (r'which part has an error', 5)]),
    ("Grammar", "Punctuations", [(r'punctuation', 5), (r'comma', 3), (r'semicolon', 4), (r'apostrophe', 4), (r'properly punctuated', 5)]),
    ("Grammar", "Parts of Speech", [(r'parts of speech', 5), (r'noun', 3), (r'verb', 3), (r'adjective', 3), (r'adverb', 3), (r'preposition', 3), (r'conjunction', 3), (r'pronoun', 3)]),
    ("Grammar", "Subject–Verb Agreement", [(r'subject-verb', 5), (r'subject verb agreement', 5), (r'singular verb', 4), (r'plural verb', 4), (r'agrees with subject', 5)]),

    # VA
    ("VA", "RC", [(r'passage', 4), (r'according to the passage', 5), (r'author implies', 5), (r'main idea of the passage', 5), (r'reading comprehension', 5), (r'passage below', 4), (r'author\'s tone', 5), (r'central theme', 5)]),
    ("VA", "Para Completion", [(r'complete the paragraph', 5), (r'para completion', 5), (r'completes the paragraph', 5), (r'sentence fits best', 4)]),
    ("VA", "Para Jumbles", [(r'para jumble', 5), (r'rearrange', 5), (r'jumbled', 5), (r'proper sequence', 5), (r'logical order', 4), (r'sentence A, B, C', 5), (r'1, 2, 3, 4', 2)]),
    ("VA", "Sentence Correction", [(r'sentence correction', 5), (r'phrase replacement', 5), (r'sentence improvement', 5), (r'replace the underlined', 5)]),
    ("VA", "Verbal Analogy", [(r'analogy', 5), (r'is to .* as', 5), (r'analogous', 5), (r'verbal analogy', 5)])
]

QUANTS_PATTERNS = [
    # Arithmetic
    ("Arithmetic", "Averages", [(r'average', 4), (r'mean age', 4), (r'weighted average', 5), (r'average score', 4)]),
    ("Arithmetic", "Mixtures & Alligation", [(r'mixture', 5), (r'alligation', 5), (r'solution contains', 4), (r'milk and water', 5), (r'replacement of liquid', 5)]),
    ("Arithmetic", "Percentages", [(r'percent', 4), (r'percentage', 4), (r'increase by \d+%', 5), (r'decrease by \d+%', 5), (r'% of', 3)]),
    ("Arithmetic", "Profit & Loss", [(r'cost price', 5), (r'selling price', 5), (r'profit', 4), (r'loss', 4), (r'discount', 4), (r'marked price', 5), (r'gain percent', 5)]),
    ("Arithmetic", "Ratio & Proportion", [(r'ratio', 4), (r'proportion', 4), (r'proportional', 4), (r'divided in the ratio', 5), (r'a : b', 4)]),
    ("Arithmetic", "SI/CI", [(r'simple interest', 5), (r'compound interest', 5), (r'principal amount', 5), (r'rate of interest', 5), (r'compounded annually', 5)]),
    ("Arithmetic", "Time & Work", [(r'time and work', 5), (r'efficiency', 3), (r'pipes and cistern', 5), (r'days to complete', 4), (r'can finish a job in', 5)]),
    ("Arithmetic", "Time–Speed–Distance", [(r'speed', 4), (r'distance', 3), (r'train', 4), (r'relative speed', 5), (r'boat and stream', 5), (r'km/h', 4), (r'm/s', 4), (r'average speed', 5)]),

    # Number System
    ("Number System", "Digit properties", [(r'unit digit', 5), (r'ten\'s digit', 5), (r'two-digit number', 5), (r'sum of digits', 4)]),
    ("Number System", "Divisibility rules", [(r'divisible by', 5), (r'divisibility', 5), (r'multiple of', 3)]),
    ("Number System", "Factorials", [(r'factorial', 5), (r'\bn!\b', 5), (r'number of zeroes at the end', 5)]),
    ("Number System", "Factorization", [(r'factorization', 5), (r'prime factor', 5), (r'prime factorization', 5)]),
    ("Number System", "Factors/Multiples", [(r'number of factors', 5), (r'sum of factors', 5), (r'number of divisors', 5)]),
    ("Number System", "HCF/LCM", [(r'hcf', 5), (r'lcm', 5), (r'greatest common divisor', 5), (r'highest common factor', 5)]),
    ("Number System", "Integral Solution", [(r'integral solution', 5), (r'integer solutions', 5), (r'positive integer solutions', 5)]),
    ("Number System", "Remainders", [(r'remainder', 5), (r'remains when divided', 5), (r'remainder theorem', 5)]),
    ("Number System", "Unit digits", [(r'unit digit of', 5), (r'last digit of', 5)]),
    ("Number System", "Miscellaneous", [(r'number system', 4), (r'real number', 3), (r'irrational number', 4)]),

    # Algebra
    ("Algebra", "Binomial Theorem", [(r'binomial', 5), (r'expansion of', 4), (r'coefficient of x\^', 5)]),
    ("Algebra", "Matrices & Determinants", [(r'matrix', 5), (r'matrices', 5), (r'determinant', 5), (r'eigenvalue', 5)]),
    ("Algebra", "Algebraic identities", [(r'algebraic', 4), (r'a\^2\s*\+\s*b\^2', 5), (r'x\s*\+\s*1/x', 5)]),
    ("Algebra", "Functions", [(r'f\(x\)', 5), (r'domain of', 4), (r'range of f', 5), (r'composite function', 5), (r'f\(f\(x\)\)', 5)]),
    ("Algebra", "Indices & Surds", [(r'indices', 5), (r'surds', 5), (r'exponent', 4), (r'power of \d+', 3)]),
    ("Algebra", "Inequalities", [(r'inequality', 5), (r'inequalities', 5), (r'greater than or equal to', 4), (r'less than or equal to', 4)]),
    ("Algebra", "Linear/Quadratic equations", [(r'quadratic equation', 5), (r'roots of the equation', 5), (r'discriminant', 5), (r'linear equation', 4)]),
    ("Algebra", "Maxima & Minima", [(r'maxima', 5), (r'minima', 5), (r'maximum value', 5), (r'minimum value', 5)]),
    ("Algebra", "Modulus", [(r'modulus', 5), (r'\|x\|', 5), (r'absolute value', 5)]),
    ("Algebra", "Polynomials", [(r'polynomial', 5), (r'degree of polynomial', 5), (r'cubic polynomial', 5)]),
    ("Algebra", "Progressions", [(r'arithmetic progression', 5), (r'geometric progression', 5), (r'harmonic progression', 5), (r'\bap\b', 3), (r'\bgp\b', 3), (r'sum of n terms', 5)]),
    ("Algebra", "Sets", [(r'subset', 5), (r'union of sets', 5), (r'intersection of sets', 5)]),

    # Geometry & Mensuration
    ("Geometry & Mensuration", "Area & Perimeter", [(r'area of', 4), (r'perimeter', 5), (r'circumference', 5)]),
    ("Geometry & Mensuration", "Circles", [(r'circle', 5), (r'radius', 4), (r'diameter', 4), (r'chord', 5), (r'tangent', 5), (r'secant', 5)]),
    ("Geometry & Mensuration", "Coordinate Geometry", [(r'coordinate', 5), (r'slope of line', 5), (r'y-intercept', 5), (r'distance formula', 5), (r'locus', 5)]),
    ("Geometry & Mensuration", "Heights & Distances", [(r'angle of elevation', 5), (r'angle of depression', 5), (r'height of tower', 5)]),
    ("Geometry & Mensuration", "Lines & Angles", [(r'parallel lines', 5), (r'transversal', 5), (r'alternate interior angle', 5)]),
    ("Geometry & Mensuration", "Polygons", [(r'polygon', 5), (r'hexagon', 5), (r'pentagon', 5), (r'diagonals of polygon', 5)]),
    ("Geometry & Mensuration", "Quadrilaterals", [(r'rectangle', 4), (r'square', 3), (r'parallelogram', 5), (r'rhombus', 5), (r'trapezium', 5), (r'quadrilateral', 5)]),
    ("Geometry & Mensuration", "Solids", [(r'sphere', 5), (r'cylinder', 5), (r'cone', 5), (r'cube', 4), (r'cuboid', 5), (r'volume of', 4), (r'surface area', 5)]),
    ("Geometry & Mensuration", "Triangles", [(r'triangle', 5), (r'hypotenuse', 5), (r'isosceles', 5), (r'equilateral', 5), (r'pythagoras', 5), (r'inradius', 5), (r'circumradius', 5)]),
    ("Geometry & Mensuration", "Trigonometry", [(r'sin', 4), (r'cos', 4), (r'tan', 4), (r'cot', 4), (r'sec', 4), (r'cosec', 4), (r'trigonometric identity', 5)]),

    # Modern Maths
    ("Modern Maths", "Logarithm", [(r'logarithm', 5), (r'\blog\b', 4), (r'log_\d+', 5)]),
    ("Modern Maths", "P & C", [(r'permutation', 5), (r'combination', 5), (r'\bncr\b', 5), (r'\bnpr\b', 5), (r'arranged in a row', 5), (r'ways of selecting', 5)]),
    ("Modern Maths", "Probability", [(r'probability', 5), (r'randomly selected', 5), (r'fair dice', 5), (r'pack of cards', 5), (r'favourable outcomes', 5)]),
    ("Modern Maths", "Set Theory", [(r'venn diagram', 5), (r'set theory', 5), (r'universal set', 5)])
]

LRDI_PATTERNS = [
    ("Logical Reasoning", "Arrangements", [(r'seating arrangement', 5), (r'circular table', 5), (r'facing north', 5), (r'linear arrangement', 5), (r'sitting in a row', 5)]),
    ("Logical Reasoning", "Blood Relations", [(r'blood relation', 5), (r'brother of', 4), (r'sister of', 4), (r'mother of', 4), (r'father of', 4), (r'paternal uncle', 5)]),
    ("Logical Reasoning", "Clocks & Calendars", [(r'clock', 5), (r'calendar', 5), (r'day of the week', 4), (r'hands of a clock', 5), (r'leap year', 5)]),
    ("Logical Reasoning", "Coding-Decoding", [(r'coding', 5), (r'decoding', 5), (r'coded as', 5), (r'written as', 3), (r'cipher', 5)]),
    ("Logical Reasoning", "Direction Sense", [(r'direction', 4), (r'walks north', 5), (r'turns right', 4), (r'turns left', 4), (r'facing south', 5)]),
    ("Logical Reasoning", "Syllogisms", [(r'syllogism', 5), (r'all .* are .*', 4), (r'some .* are not', 4), (r'no .* is a', 4), (r'conclusion follows', 5)]),
    ("Logical Reasoning", "Series & Analogies", [(r'number series', 5), (r'missing number', 5), (r'find the next number', 5), (r'pattern', 3)]),
    ("Logical Reasoning", "Venn Diagrams", [(r'venn diagram', 5), (r'represented by circle', 4), (r'overlapping regions', 5)]),

    ("Data Interpretation", "Bar Charts", [(r'bar chart', 5), (r'bar graph', 5), (r'horizontal bars', 5)]),
    ("Data Interpretation", "Line Graphs", [(r'line graph', 5), (r'line chart', 5), (r'trend over years', 4)]),
    ("Data Interpretation", "Pie Charts", [(r'pie chart', 5), (r'degrees in pie chart', 5), (r'sector angle', 5)]),
    ("Data Interpretation", "Tables", [(r'table shows', 5), (r'refer to the table', 5), (r'tabular data', 5)]),
    ("Data Interpretation", "Caselets", [(r'caselet', 5), (r'passage based di', 5), (r'data given below', 3)]),
    ("Data Interpretation", "Data Sufficiency", [(r'data sufficiency', 5), (r'statement 1 alone', 5), (r'statement 2 alone', 5), (r'is sufficient to answer', 5)])
]

def auto_classify_topic_subtopic(question_text: str, hint_text: str = "", subject: str = "English"):
    """Classifies a question into topic/subtopic using a weighted keyword scoring algorithm (0 API Cost)."""
    combined = (question_text + " " + (hint_text or "")).lower()
    
    subj_clean = (subject or "").strip().lower()
    if "quant" in subj_clean or "math" in subj_clean:
        patterns = QUANTS_PATTERNS
        default_topic, default_subtopic = "Arithmetic", "Percentages"
    elif "lr" in subj_clean or "di" in subj_clean or "reasoning" in subj_clean:
        patterns = LRDI_PATTERNS
        default_topic, default_subtopic = "Logical Reasoning", "Arrangements"
    else:
        patterns = ENGLISH_PATTERNS
        default_topic, default_subtopic = "VA", "RC"
    
    best_score = 0
    best_candidate = (default_topic, default_subtopic)

    for topic, subtopic, pats in patterns:
        score = 0
        for pat, weight in pats:
            matches = len(re.findall(pat, combined))
            if matches > 0:
                score += weight * min(matches, 3)
        
        if score > best_score:
            best_score = score
            best_candidate = (topic, subtopic)
            
    return best_candidate

def parse_with_abacus_fallback(raw_text: str, subject: str = "English", abacus_key: str = None, model: str = "gpt-4o") -> list:
    """Uses Abacus.AI API to extract questions, options, hints, and topics from document text in chunks."""
    key = abacus_key if abacus_key and abacus_key.strip() else os.environ.get("ABACUS_API_KEY", "")
    if not key or not key.strip():
        return []
    
    import requests

    tax = OFFICIAL_TAXONOMY.get(subject, OFFICIAL_TAXONOMY.get("English", {}))
    tax_str = json.dumps(tax, indent=2)

    # Split text into chunks of ~5000 chars to avoid max_token output truncation
    chunk_size = 5000
    paragraphs = raw_text.split("\n\n")
    chunks = []
    curr = ""
    for p in paragraphs:
        if len(curr) + len(p) > chunk_size and curr.strip():
            chunks.append(curr.strip())
            curr = p
        else:
            curr += "\n\n" + p if curr else p
    if curr.strip():
        chunks.append(curr.strip())

    if not chunks:
        chunks = [raw_text[:16000]]

    all_parsed_questions = []

    for chunk in chunks:
        prompt = f"""You are an expert academic document parser and tutor. Extract all multiple choice questions from the given document snippet into a clean JSON array.

For each question:
1. Extract the full questionText (format mathematical symbols using inline LaTeX $ ... $).
2. Extract all options into an array of objects: {{"text": "...", "isCorrect": true/false}}. Identify the correct option if an answer key or solution is present.
3. Extract or generate a clear step-by-step hint/explanation for the question in the "hint" field.
4. Classify the question into the exact "topic" and "subtopic" from this official taxonomy:
Taxonomy JSON:
{tax_str}

Return ONLY a valid JSON array of objects with this exact structure:
[
  {{
    "questionText": "Full text of the question",
    "hint": "Step-by-step hint, solution, or explanation",
    "topic": "TopicName from taxonomy",
    "subtopic": "SubtopicName from taxonomy",
    "options": [
      {{"text": "Option A text", "isCorrect": false}},
      {{"text": "Option B text", "isCorrect": true}},
      {{"text": "Option C text", "isCorrect": false}},
      {{"text": "Option D text", "isCorrect": false}}
    ]
  }}
]

Strict Rules:
- Extract every question present in the snippet.
- Do not add extra conversational text outside the JSON code block.

Subject: {subject}

Document Snippet:
{chunk}"""

        model_name = model if model and model.strip() else "gpt-4o"
        routellm_endpoints = [
            "https://routellm.abacus.ai/v1/chat/completions",
            "https://paas.abacus.ai/v1/chat/completions",
            "https://api.abacus.ai/v1/chat/completions"
        ]
        parsed_chunk = False
        for ep in routellm_endpoints:
            if parsed_chunk: break
            headers = {
                "Authorization": f"Bearer {key.strip()}",
                "Content-Type": "application/json"
            }
            payload = {
                "model": model_name,
                "messages": [{"role": "user", "content": prompt}],
                "temperature": 0.1,
                "max_tokens": 4096
            }
            try:
                r = requests.post(ep, json=payload, headers=headers, timeout=30)
                if r.status_code == 200:
                    res_data = r.json()
                    choices = res_data.get("choices", [])
                    if choices:
                        txt = choices[0].get("message", {}).get("content", "").strip()
                        match = re.search(r'```(?:json)?\s*(\[.*\])\s*```', txt, re.DOTALL) or re.search(r'\[.*\]', txt, re.DOTALL)
                        json_str = match.group(1) if (match and match.lastindex) else (match.group(0) if match else txt)
                        try:
                            parsed = json.loads(json_str)
                        except Exception:
                            fixed_json = re.sub(r',\s*([\]}])', r'\1', json_str)
                            if not fixed_json.rstrip().endswith("]"):
                                fixed_json = fixed_json.rstrip() + "]"
                            parsed = json.loads(fixed_json)

                        if isinstance(parsed, list) and len(parsed) > 0:
                            for q in parsed:
                                q_text = format_math_latex(q.get("questionText", ""))
                                h_text = format_math_latex(q.get("hint", ""))
                                t_top, t_sub = auto_classify_topic_subtopic(q_text, h_text, subject=subject)
                                
                                opts = []
                                for o in q.get("options", []):
                                    opt_txt = o.get("text", "") if isinstance(o, dict) else str(o)
                                    is_corr = bool(o.get("isCorrect", False)) if isinstance(o, dict) else False
                                    opts.append({
                                        "text": format_math_latex(opt_txt),
                                        "isCorrect": is_corr
                                    })
                                
                                all_parsed_questions.append({
                                    "questionText": q_text,
                                    "hint": h_text,
                                    "topic": q.get("topic") or t_top,
                                    "subtopic": q.get("subtopic") or t_sub,
                                    "options": opts
                                })
                            parsed_chunk = True
            except Exception as e:
                print(f"Abacus AI parsing error ({ep}): {e}")

    return all_parsed_questions

def parse_with_gemini_fallback(raw_text: str, subject: str = "English", api_key: str = None) -> list:
    """Uses Google Gemini API to extract questions, options, hints, and topics from document text in chunks."""
    key = api_key if api_key and api_key.strip() else os.environ.get("GEMINI_API_KEY", "")
    if not key or not key.strip():
        return []
    
    import requests

    tax = OFFICIAL_TAXONOMY.get(subject, OFFICIAL_TAXONOMY.get("English", {}))
    tax_str = json.dumps(tax, indent=2)

    chunk_size = 5000
    paragraphs = raw_text.split("\n\n")
    chunks = []
    curr = ""
    for p in paragraphs:
        if len(curr) + len(p) > chunk_size and curr.strip():
            chunks.append(curr.strip())
            curr = p
        else:
            curr += "\n\n" + p if curr else p
    if curr.strip():
        chunks.append(curr.strip())

    if not chunks:
        chunks = [raw_text[:16000]]

    all_parsed_questions = []

    models = ["gemini-2.0-flash", "gemini-1.5-flash-latest", "gemini-1.5-flash"]
    for chunk in chunks:
        prompt = f"""You are an expert academic document parser and tutor. Extract all multiple choice questions from the given document text into a clean JSON array.

For each question:
1. Extract the full questionText (format mathematical symbols using inline LaTeX $ ... $).
2. Extract all options into an array of objects: {{"text": "...", "isCorrect": true/false}}. Identify the correct option if an answer key or solution is present.
3. Extract or generate a clear step-by-step hint/explanation for the question in the "hint" field.
4. Classify the question into the exact "topic" and "subtopic" from this official taxonomy:
Taxonomy JSON:
{tax_str}

Return ONLY a valid JSON array of objects with this exact structure:
[
  {{
    "questionText": "Full text of the question",
    "hint": "Step-by-step hint, solution, or explanation",
    "topic": "TopicName from taxonomy",
    "subtopic": "SubtopicName from taxonomy",
    "options": [
      {{"text": "Option A text", "isCorrect": false}},
      {{"text": "Option B text", "isCorrect": true}},
      {{"text": "Option C text", "isCorrect": false}},
      {{"text": "Option D text", "isCorrect": false}}
    ]
  }}
]

Strict Rules:
- Extract every question present in the text.
- Do not add extra conversational text outside the JSON code block.

Subject: {subject}

Document Text:
{chunk}"""

        payload = {
            "contents": [{"parts": [{"text": prompt}]}],
            "generationConfig": {"temperature": 0.1, "maxOutputTokens": 4096}
        }
        
        parsed_chunk = False
        for m in models:
            if parsed_chunk: break
            url = f"https://generativelanguage.googleapis.com/v1beta/models/{m}:generateContent?key={key.strip()}"
            try:
                r = requests.post(url, json=payload, headers={"Content-Type": "application/json"}, timeout=25)
                if r.status_code == 200:
                    res_data = r.json()
                    candidates = res_data.get('candidates', [])
                    if candidates:
                        cand = candidates[0]
                        content_obj = cand.get('content')
                        if content_obj and isinstance(content_obj, dict):
                            parts = content_obj.get('parts', [])
                            txt = "\n".join([p.get('text', '') for p in parts if isinstance(p, dict) and 'text' in p]).strip()
                            match = re.search(r'```(?:json)?\s*(\[.*\])\s*```', txt, re.DOTALL) or re.search(r'\[.*\]', txt, re.DOTALL)
                            json_str = match.group(1) if (match and match.lastindex) else (match.group(0) if match else txt)
                            try:
                                parsed = json.loads(json_str)
                            except Exception:
                                fixed_json = re.sub(r',\s*([\]}])', r'\1', json_str)
                                if not fixed_json.rstrip().endswith("]"):
                                    fixed_json = fixed_json.rstrip() + "]"
                                parsed = json.loads(fixed_json)

                            if isinstance(parsed, list) and len(parsed) > 0:
                                for q in parsed:
                                    q_text = format_math_latex(q.get("questionText", ""))
                                    h_text = format_math_latex(q.get("hint", ""))
                                    t_top, t_sub = auto_classify_topic_subtopic(q_text, h_text, subject=subject)
                                    
                                    opts = []
                                    for o in q.get("options", []):
                                        opt_txt = o.get("text", "") if isinstance(o, dict) else str(o)
                                        is_corr = bool(o.get("isCorrect", False)) if isinstance(o, dict) else False
                                        opts.append({
                                            "text": format_math_latex(opt_txt),
                                            "isCorrect": is_corr
                                        })
                                    
                                    all_parsed_questions.append({
                                        "questionText": q_text,
                                        "hint": h_text,
                                        "topic": q.get("topic") or t_top,
                                        "subtopic": q.get("subtopic") or t_sub,
                                        "options": opts
                                    })
                                parsed_chunk = True
            except Exception as e:
                print(f"Gemini AI parsing failed with model {m}: {e}")

    return all_parsed_questions

def classify_topics_with_ai(questions: list, subject: str = "English", api_key: str = None, abacus_key: str = None, provider: str = "gemini", model: str = "gpt-4o") -> list:
    """Uses LLM to categorize topics & subtopics for parsed questions when use_ai_topics is enabled."""
    if not questions:
        return questions
        
    key = (api_key or os.environ.get("GEMINI_API_KEY", "")).strip()
    ab_key = (abacus_key or os.environ.get("ABACUS_API_KEY", "")).strip()
    
    if not key and not ab_key:
        return questions

    tax = OFFICIAL_TAXONOMY.get(subject, OFFICIAL_TAXONOMY.get("English", {}))
    tax_str = json.dumps(tax, indent=2)

    import requests
    batch_size = 15
    for i in range(0, len(questions), batch_size):
        batch = questions[i:i+batch_size]
        q_summary = []
        for idx, q in enumerate(batch):
            opts_formatted = []
            for j, o in enumerate(q.get('options', [])):
                opt_str = o.get('text', '') if isinstance(o, dict) else str(o)
                opts_formatted.append(f"{chr(65+j)}: {opt_str}")
            opts = ", ".join(opts_formatted)
            q_text = (q.get('questionText', '') or '')[:300]
            q_summary.append(f"Q{idx+1}: {q_text} (Options: {opts})")
        
        prompt = f"""Classify each of the following questions into the exact topic and subtopic from this official taxonomy:
Taxonomy JSON:
{tax_str}

Subject: {subject}

Questions:
""" + "\n\n".join(q_summary) + """

Return ONLY a valid JSON array of objects with the exact classification:
[
  {{"qIndex": 1, "topic": "TopicName", "subtopic": "SubtopicName"}}
]"""

        try:
            if provider.lower() == "abacus" or (ab_key and not key):
                headers = {"Authorization": f"Bearer {ab_key}", "Content-Type": "application/json"}
                ep = "https://routellm.abacus.ai/v1/chat/completions"
                res = requests.post(ep, json={"model": model, "messages": [{"role": "user", "content": prompt}], "temperature": 0.1}, headers=headers, timeout=20)
                if res.status_code == 200:
                    txt = res.json().get("choices", [{}])[0].get("message", {}).get("content", "")
                    match = re.search(r'```(?:json)?\s*(\[.*\])\s*```', txt, re.DOTALL) or re.search(r'\[.*\]', txt, re.DOTALL)
                    if match:
                        json_str = match.group(1) if (match and match.lastindex) else match.group(0)
                        cls_list = json.loads(json_str)
                        for item in cls_list:
                            q_idx = item.get("qIndex", 0) - 1
                            if 0 <= q_idx < len(batch):
                                if item.get("topic"): batch[q_idx]["topic"] = item["topic"]
                                if item.get("subtopic"): batch[q_idx]["subtopic"] = item["subtopic"]
            elif key:
                gemini_models = ["gemini-2.0-flash", "gemini-1.5-flash-latest", "gemini-1.5-flash"]
                for g_model in gemini_models:
                    url = f"https://generativelanguage.googleapis.com/v1beta/models/{g_model}:generateContent?key={key}"
                    res = requests.post(url, json={"contents": [{"parts": [{"text": prompt}]}], "generationConfig": {"temperature": 0.1}}, headers={"Content-Type": "application/json"}, timeout=20)
                    if res.status_code == 200:
                        candidates = res.json().get('candidates', [])
                        if candidates:
                            cand = candidates[0]
                            content_obj = cand.get('content')
                            if content_obj and isinstance(content_obj, dict):
                                parts = content_obj.get('parts', [])
                                txt = "\n".join([p.get('text', '') for p in parts if isinstance(p, dict) and 'text' in p]).strip()
                                match = re.search(r'```(?:json)?\s*(\[.*\])\s*```', txt, re.DOTALL) or re.search(r'\[.*\]', txt, re.DOTALL)
                                if match:
                                    json_str = match.group(1) if (match and match.lastindex) else match.group(0)
                                    try:
                                        cls_list = json.loads(json_str)
                                    except Exception:
                                        fixed_json = re.sub(r',\s*([\]}])', r'\1', json_str)
                                        cls_list = json.loads(fixed_json)
                                    if isinstance(cls_list, list):
                                        for item in cls_list:
                                            q_idx = item.get("qIndex", 0) - 1
                                            if 0 <= q_idx < len(batch):
                                                if item.get("topic"): batch[q_idx]["topic"] = item["topic"]
                                                if item.get("subtopic"): batch[q_idx]["subtopic"] = item["subtopic"]
                                        break
        except Exception as e:
            print(f"AI Topic classification batch warning: {e}")

    return questions

def parse_pdf_questions(
    file_path: str,
    subject: str = "English",
    default_topic: str = None,
    default_subtopic: str = None,
    api_key: str = None,
    abacus_key: str = None,
    provider: str = "gemini",
    model: str = "gpt-4o",
    use_ai_topics: bool = False,
    custom_prompt: str = None
) -> list:
    """
    Parses a question paper PDF or DOCX into JSON.
    When API keys are available, Stage 1 uses AI-only extraction directly from the PDF text.
    If no API key is provided or AI extraction returns 0 questions, local regex parser runs as fallback.
    """
    text = extract_raw_text(file_path)
    if not text or not text.strip():
        return []

    provider_clean = (provider or "gemini").lower().strip()
    has_gemini_key = bool(api_key and api_key.strip()) or bool(os.environ.get("GEMINI_API_KEY", "").strip())
    has_abacus_key = bool(abacus_key and abacus_key.strip()) or bool(os.environ.get("ABACUS_API_KEY", "").strip())

    # =========================================================================
    # STAGE 1: AI-ONLY EXTRACTION (PRIMARY STAGE WHEN API KEYS ARE AVAILABLE)
    # =========================================================================
    if has_gemini_key or has_abacus_key:
        ai_questions = []
        if provider_clean == "abacus" or (has_abacus_key and not has_gemini_key):
            ai_questions = parse_with_abacus_fallback(text, subject=subject, abacus_key=abacus_key, model=model)
            if not ai_questions and has_gemini_key:
                ai_questions = parse_with_gemini_fallback(text, subject=subject, api_key=api_key)
        else:
            ai_questions = parse_with_gemini_fallback(text, subject=subject, api_key=api_key)
            if not ai_questions and has_abacus_key:
                ai_questions = parse_with_abacus_fallback(text, subject=subject, abacus_key=abacus_key, model=model)

        if ai_questions and len(ai_questions) > 0:
            if use_ai_topics:
                ai_questions = classify_topics_with_ai(ai_questions, subject=subject, api_key=api_key, abacus_key=abacus_key, provider=provider, model=model)
            return ai_questions

    # =========================================================================
    # STAGE 2: LOCAL REGEX PARSER (FALLBACK WHEN NO API KEY OR AI FINDS 0 QUESTIONS)
    # =========================================================================
    q_marker_pattern = r'(?:^|\n)\s*(?:Q(?:uestion)?\s*[-.]?\s*\d+[:\.-]?|\b\d+[\.\):-]|(?:\([0-9]+\)))\s*'
    matches = list(re.finditer(q_marker_pattern, text, re.MULTILINE | re.IGNORECASE))

    question_blocks = []
    for i in range(len(matches)):
        start_idx = matches[i].start()
        end_idx = matches[i+1].start() if i + 1 < len(matches) else len(text)
        block_text = text[start_idx:end_idx].strip()
        question_blocks.append(block_text)

    parsed_questions = []

    for block in question_blocks:
        marker_match = re.match(r'^(Q(?:uestion)?\s*[-.]?\s*\d+[:\.-]?|\d+[\.\):-]|(?:\([0-9]+\)))\s*', block, re.IGNORECASE)
        body = block
        if marker_match:
            body = block[marker_match.end():].strip()

        sol_parts = []
        correct_letter = ""

        ans_inline_match = re.search(r'\(?Ans(?:wer)?:\s*([A-E0-9a-zA-Z]+)\)?', body, re.IGNORECASE)
        if ans_inline_match:
            ans_val = ans_inline_match.group(1).strip().upper()
            if ans_val in ['A', 'B', 'C', 'D', 'E']:
                correct_letter = ans_val
            sol_parts.append(f"Correct Answer: {ans_val}")
            body = body[:ans_inline_match.start()].strip() + "\n" + body[ans_inline_match.end():].strip()
        
        explanation_parts = []
        
        exp_match = re.search(r'(?:Explanation|Solution|Sol|Hint|Reason|Rationale):\s*(.*?)(?=(?:Why the other options are wrong:|Correct Answer:|Answer:|Key:|$))', body, re.DOTALL | re.IGNORECASE)
        if exp_match:
            exp_text = exp_match.group(1).strip()
            if exp_text and len(exp_text) > 2:
                explanation_parts.append(exp_text)

        why_wrong_match = re.search(r'Why the other options are wrong:?\s*(.*?)(?=(?:Q\s*\d+[:\.]|Question\s+\d+[:\.]|\b\d+[\.\)]|$))', body, re.DOTALL | re.IGNORECASE)
        if why_wrong_match:
            why_text = why_wrong_match.group(1).strip()
            if why_text and len(why_text) > 2:
                explanation_parts.append("Why the other options are wrong:\n" + why_text)

        clean_body = body
        if why_wrong_match:
            clean_body = clean_body[:why_wrong_match.start()].strip()
        if exp_match:
            clean_body = clean_body[:exp_match.start()].strip()

        if not correct_letter:
            ans_match = re.search(r'(?:Correct Answer|Answer|Key):\s*([A-E])\)?\s*(.*)', clean_body, re.IGNORECASE)
            if ans_match:
                correct_letter = ans_match.group(1).upper()
                ans_extra = ans_match.group(2).strip()
                if ans_extra:
                    explanation_parts.append(ans_extra)
                clean_body = clean_body[:ans_match.start()].strip()

        opt_start_match = re.search(r'(?:^|\n|\s{2,})(?:([A-Ea-e1-5])[\.\):\:-]|\(([A-Ea-e1-5])\))\s*', clean_body)
        
        if opt_start_match:
            question_text = clean_body[:opt_start_match.start()].strip()
            options_text = clean_body[opt_start_match.start():].strip()
        else:
            question_text = clean_body.strip()
            options_text = ""

        raw_options = []
        if options_text:
            opt_matches = list(re.finditer(r'(?:^|\n|\s{2,})(?:([A-Ea-e1-5])[\.\):\:-]|\(([A-Ea-e1-5])\))', options_text))
            for idx in range(len(opt_matches)):
                op_m = opt_matches[idx]
                letter = (op_m.group(1) or op_m.group(2)).upper()
                if letter.isdigit():
                    letter_num = int(letter)
                    if 1 <= letter_num <= 5:
                        letter = chr(64 + letter_num)
                start_op = op_m.end()
                end_op = opt_matches[idx+1].start() if idx + 1 < len(opt_matches) else len(options_text)
                opt_val = options_text[start_op:end_op].strip()
                raw_options.append((letter, opt_val))

        formatted_options = []
        for letter, opt_val in raw_options:
            is_corr = False
            if correct_letter:
                is_corr = (letter == correct_letter)

            formatted_options.append({
                "text": format_math_latex(opt_val),
                "isCorrect": is_corr
            })

        if correct_letter and formatted_options:
            for opt, (letter, _) in zip(formatted_options, raw_options):
                if letter == correct_letter:
                    opt["isCorrect"] = True

        if explanation_parts:
            hint = "\n\n".join(explanation_parts)
        else:
            hint = ""

        q_num_str = str(len(parsed_questions) + 1)
        if marker_match:
            digits = re.search(r'\d+', marker_match.group(0))
            if digits:
                q_num_str = digits.group(0)

        if len(question_text) > 120 and ("passage" in question_text.lower() or len(question_text.splitlines()) > 3):
            last_passage = question_text
        elif not question_text:
            if 'last_passage' in locals() and last_passage:
                question_text = f"{last_passage}\n\nQuestion {q_num_str}"
            else:
                question_text = f"Question {q_num_str}"

        question_text = format_math_latex(question_text)
        hint = format_math_latex(hint)

        if question_text or formatted_options:
            auto_top, auto_sub = auto_classify_topic_subtopic(question_text, hint, subject=subject)
            parsed_questions.append({
                "questionText": question_text,
                "hint": hint,
                "topic": default_topic or auto_top,
                "subtopic": default_subtopic or auto_sub,
                "options": formatted_options
            })

    if parsed_questions and len(parsed_questions) > 0:
        if use_ai_topics:
            parsed_questions = classify_topics_with_ai(parsed_questions, subject=subject, api_key=api_key, abacus_key=abacus_key, provider=provider, model=model)
        return parsed_questions

    return []



